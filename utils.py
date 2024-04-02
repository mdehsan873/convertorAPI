import io
import os
import fitz
import os
from docx import Document
from PIL import Image


def word_to_pdf(word_path, pdf_path):
    doc = Document(word_path)
    pdf = fitz.open()
    for paragraph in doc.paragraphs:
        pdf_page = pdf.new_page()
        pdf_page.insert_text((10, 10), paragraph.text)
    pdf.save(pdf_path)


def image_to_pdf(image_paths, pdf_path):
    pdf = fitz.open()
    for image_path in image_paths:
        img = Image.open(image_path)
        img = fitz.open(stream=img.tobytes())
        pdf_page = pdf.new_page(width=img.width, height=img.height)
        pdf_page.show_image(0, fitz.Rect(0, 0, img.width, img.height), stream=img.convert("RGB").tobytes())
    pdf.save(pdf_path)


def pdf_to_word(pdf_path, word_path):
    doc = fitz.open(stream=io.BytesIO(pdf_path))
    output = ""
    paths = []
    for page in doc:
        output += page.get_text()

    # Create a new Word document
    docx_document = Document()

    # Add the text to the Word document
    docx_document.add_paragraph(output)

    # Save the Word document as .docx format compatible with Word 2007
    docx_document.save(word_path)
    paths.append(word_path)
    return paths


def pdf_to_images(pdf_path, image_folder):
    doc = fitz.open(stream=io.BytesIO(pdf_path))
    image_paths = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap()
        image_path = os.path.join(image_folder, f"page_{i}.png")
        pix.save(image_path)
        image_paths.append(image_path)
    return image_paths


def pdf_to_jpg(pdf_path, image_folder):
    doc = fitz.open(stream=io.BytesIO(pdf_path))
    image_paths = []
    for i, page in enumerate(doc):
        pix = page.get_pixmap()
        image_path = os.path.join(image_folder, f"page_{i}.jpg")
        pix.save(image_path)
        image_paths.append(image_path)
    return image_paths


def word_to_pdf(word_path, pdf_path):
    doc = Document(io.BytesIO(word_path))
    pdf = fitz.open()
    for paragraph in doc.paragraphs:
        pdf_page = pdf.new_page()
        pdf_page.insert_text((10, 10), paragraph.text)
    pdf.save('pdf_path.pdf')
    return [pdf_path]


def image_to_pdf(image_paths, pdf_path):
    pdf = fitz.open()
    img = fitz.open(stream=io.BytesIO(image_paths))
    img_pixmap = img[0].get_pixmap()
    pdf_page = pdf.new_page(width=img_pixmap.width, height=img_pixmap.height)
    pdf_page.insert_image(fitz.Rect(0, 0, img_pixmap.width, img_pixmap.height), pixmap=img_pixmap)
    img.close()
    pdf.save(pdf_path)
    pdf.close()
    return [pdf_path]


if __name__ == "__main__":
    # pdf_file = "Ehsan_Backend.pdf"
    # word_output = "output.docx"
    # image_output_folder = "images"
    #
    # # Convert PDF to Word
    # pdf_to_word(pdf_file, word_output)
    # print("PDF converted to Word.")
    #
    # # Convert PDF to PNG images
    # if not os.path.exists(image_output_folder):
    #     os.makedirs(image_output_folder)
    # pdf_to_images(pdf_file, image_output_folder)
    # print("PDF converted to PNG images.")
    #
    # # Convert PDF to JPG images
    # if not os.path.exists(image_output_folder):
    #     os.makedirs(image_output_folder)
    # pdf_to_jpg(pdf_file, image_output_folder)
    # print("PDF converted to JPG images.")
    word_file = "/Users/ehsanansari/PycharmProjects/eagle-tft/output.docx"
    word_to_pdf(word_file, "output_word_to_pdf.pdf")
    print("Word converted to PDF.")

    # Convert PNG images to PDF
    # png_files = ["image1.png", "image2.png"]  # Add paths to your PNG files
    # image_to_pdf(png_files, "output_png_to_pdf.pdf")
    # print("PNG images converted to PDF.")
    #
    # # Convert JPG images to PDF
    # jpg_files = ["image1.jpg", "image2.jpg"]  # Add paths to your JPG files
    # image_to_pdf(jpg_files, "output_jpg_to_pdf.pdf")
    # print("JPG images converted to PDF.")
